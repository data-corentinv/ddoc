import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import scipy.stats as ss
import scipy as sp
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdb
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import roc_auc_score


class Addon:
    """Classe permettant de gérer des addons pour le document word"""


    def __call__(self, document, metadata, df):
        """
        Doit retourner document avec les modifications opérées.
        Les metadata est un dictionnaire contenant le champ
        """
        raise NotImplementedError


class FieldAddon:
    """Classe permettant de gérer des addons pour le document word"""
    def __init__(self,  field):
        self.field = field

    def __call__(self, document, metadata, series):
        """
        Doit retourner document avec les modifications opérées.
        Les metadata est un dictionnaire contenant le champ
        """
        raise NotImplementedError

class ExampleFieldAddon(FieldAddon):
    def __call__(self, document, metadata, df, field):
        series = df[field]
        p = document.add_paragraph()
        p.add_run("Ceci est un champ de taille ")
        p.add_run(str(series.shape))
        return document

class ExampleAddon(Addon):
    def __call__(self, document, metadata, df):
        p = document.add_paragraph()
        p.add_run("Terminons par cette phrase")
        p.add_run(str(df.columns))
        return document

def cat2catcorrelation(df, name_var1, name_var2):
    # Destroy intersection of nan
    df_copy = df[(df[name_var1].notnull()) | (df[name_var2].notnull())]
    crosstab = pd.crosstab(df_copy[name_var1].map(str), df_copy[name_var2].map(str), rownames=[name_var1], colnames=[name_var2])
    chi2 = ss.chi2_contingency(crosstab)[0]
    n = sum(crosstab.sum())
    return np.sqrt(chi2 / (n*(min(crosstab.shape)-1+1e-12)))


def getRFcorrelation(df, name_var1, name_var2, model, train_size = 0.7, target = "label", pos_label =True):

    #df = df.sort_index(axis=1)
    # Split
    df['__index__'] = np.arange(0,df.shape[0])
    df['__training__'] = 0
    df['__testing__'] = 0
    df.loc[df['__index__']<train_size*df.shape[0],'__training__' ] = 1
    df.loc[df['__index__']>=train_size*df.shape[0],'__testing__' ] = 1
    df[target] = df['TARGET'].apply(lambda x : x!='0.0')

    df_train = df[df['__training__']==1]
    df_test = df[df['__testing__']==1]



    # Dataset creation
    y_train = df_train[target].values
    x_train = df_train[[name_var1, name_var2]].values
    y_test = df_test[target].values
    x_test = df_test[[name_var1, name_var2]].values

    # Training & prediction
    model.fit(x_train,y_train)

    y_pred = model.predict(x_train)
    df_train['__y_pred__'] = y_pred
    try:
        y_proba = pd.DataFrame(model.predict_proba(x_train))[1]
    except:
        y_proba = y_pred
    #pdb.set_trace()
    df_train['__y_proba__'] = y_proba

    y_pred = model.predict(x_test)
    df_test['__y_pred__'] = y_pred
    try:
        y_proba = pd.DataFrame(model.predict_proba(x_test))[1]
    except:
        y_proba = y_pred
    df_test['__y_proba__'] = y_proba

    df = pd.concat([df_train, df_test], axis=0)
    df = df.sort_values(by='__index__')




    indicator ='__testing__'
    sourced =df.query(str(indicator)+"==1")
    y_true = sourced[target].values == pos_label
    y_pred = sourced['__y_pred__'].values == pos_label
    result = roc_auc_score(y_true, y_score=y_proba)

    return result





def if_else(condition, a, b) :
   if condition : return a
   else         : return b


class CorrelationAddon(Addon):

    def __call__(self, document, metadata, df):
        document.add_heading('Analyse des corrélations', 2)

        p = document.add_paragraph()
        p.add_run("Affichage des corrélations")



        all_variables = pd.DataFrame(df.columns, columns = ['variable'])

        types = []
        sous_types = []
        status = []
        for var in all_variables['variable']:
            types.append(metadata.get(str(var), {}).get('type', 'inconnu'))
            sous_types.append(metadata.get(str(var), {}).get('sous-type', 'inconnu'))
            status.append(metadata.get(str(var), {}).get('alerte', 'inconnu'))
        types = pd.DataFrame(types, columns = ['type'])
        sous_types = pd.DataFrame(types, columns = ['sous_type'])
        status = pd.DataFrame(status, columns = ['status'])

        all_variables = pd.concat([all_variables, types, status, sous_types], axis = 1)
        all_variables = all_variables.set_index(all_variables['variable'])
        all_variables = all_variables.query(
            "(status == 'modelisable' or status == 'utilisable') and (type == 'categoriel')")

        corr_df = pd.DataFrame(np.zeros((all_variables.shape[0], all_variables.shape[0])),index=all_variables['variable'], columns=all_variables['variable'])

        for num_var1, name_var1 in enumerate(all_variables['variable']):

            type_var1 = all_variables['type'][name_var1]

            if (type_var1 == 'categoriel'):
                if (df[name_var1].isnull().apply(lambda x : 1 if x else 0).sum()>0 ):
                    df[name_var1] = df[name_var1].cat.add_categories(['aaValuezz'])
                    df[name_var1] = df[name_var1].fillna(value = 'aaValuezz')


            for num_var2, name_var2 in enumerate(all_variables['variable']):
                if num_var2 <= num_var1:
                    continue
                print(str(num_var1)+':'+str(num_var2)+' ; '+ name_var1+':'+name_var2)
                type_var2 = all_variables['type'][name_var2]
                if (type_var2 == 'categoriel'):
                    if (df[name_var2].isnull().apply(lambda x : 1 if x else 0).sum()>0 ):
                        df[name_var2] = df[name_var2].cat.add_categories(['aaValuezz'])
                        df[name_var2] = df[name_var2].fillna(value = 'aaValuezz')


                if ((type_var1 == 'categoriel') and (type_var2 == 'categoriel')):
                    result = cat2catcorrelation(df, name_var1, name_var2)
                else:
                    result = 0.0

                corr_df[name_var1][num_var2] = result
                corr_df[name_var2][num_var1] = result

        # Handle the errors when there is only one modality in a variable
        corr_df = corr_df.fillna(value = 0)
        for num_var1, name_var1 in enumerate(all_variables['variable']):
            corr_df[name_var1][num_var1] = 1.0

        # # Set up the matplotlib figure
        # from matplotlib import rcParams
        # rcParams.update({'figure.autolayout': True})
        # fig, ax = plt.subplots()
        # name = "temp.png"
        # # Draw the heatmap using seaborn
        # #heatmap = sns.heatmap(corr_df, vmax=1.0, square=True, cmap=sns.diverging_palette(10, 220, as_cmap=True))
        # heatmap = sns.heatmap(corr_df, vmax=1.0, square=True, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))
        # plt.xticks(rotation = 'vertical')
        # plt.yticks(rotation = 'horizontal')
        # heatmap.get_figure().savefig(name)
        # plt.gcf().clear()
        # plt.close(fig)
        # document.add_picture(name, width=Cm(17.0), height=Cm(22.0))
        # last_paragraph = document.paragraphs[-1]
        # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        linkage = sp.cluster.hierarchy.linkage(sp.spatial.distance.squareform(1-corr_df), method='average')
        fig, ax = plt.subplots()
        heatmap_clustered = sns.clustermap(corr_df, row_linkage=linkage, col_linkage=linkage, cmap=sns.light_palette("seagreen", as_cmap = True))
        #heatmap_clustered = sns.clustermap(corr_df, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))
        name = "temp_matrix.png"
        plt.setp(heatmap_clustered.ax_heatmap.get_yticklabels(), rotation=0)
        plt.setp(heatmap_clustered.ax_heatmap.get_xticklabels(), rotation=90)
        heatmap_clustered.savefig(name)
        document.add_paragraph("V de Cramer calculé pour toutes les paires de variables catégorielles. Les variables sont réordonnées par proximité.")
        #document.add_page_break()

        document.add_picture(name, width=Cm(15.0), height=Cm(15.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        #linkage = sp.cluster.hierarchy.linkage(sp.spatial.distance.squareform(1-corr_df), method='average')
        #fig, ax = plt.subplots()
        #heatmap_clustered = sns.clustermap(corr_df, row_linkage=linkage, col_linkage=linkage, row_cluster=False, col_cluster=False, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))
        ## heatmap_clustered = sns.clustermap(corr_df, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))
        #name = "temp.png"
        #plt.setp(heatmap_clustered.ax_heatmap.get_yticklabels(), rotation=0)
        #plt.setp(heatmap_clustered.ax_heatmap.get_xticklabels(), rotation=90)
        #heatmap_clustered.savefig(name)
        #document.add_paragraph("V de Cramer calculé pour toutes les paires de variables catégorielles.")
        #document.add_picture(name, width=Cm(17.0), height=Cm(22.0))
        #last_paragraph = document.paragraphs[-1]
        #last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return document



class RandomForestCorrelationAddon(Addon):

    def __call__(self, document, metadata, df):
        document.add_heading("Performance d'un randomForest", 2)

        p = document.add_paragraph()
        p.add_run("RF")



        all_variables = pd.DataFrame(df.columns, columns = ['variable'])
        types = []
        sous_types = []
        status = []
        for var in all_variables['variable']:
            types.append(metadata.get(str(var), {}).get('type', 'inconnu'))
            sous_types.append(metadata.get(str(var), {}).get('sous-type', 'inconnu'))
            status.append(metadata.get(str(var), {}).get('alerte', 'inconnu'))
        types = pd.DataFrame(types, columns = ['type'])
        sous_types = pd.DataFrame(types, columns = ['sous_type'])
        status = pd.DataFrame(status, columns = ['status'])


        all_variables = pd.concat([all_variables, types, status, sous_types], axis = 1)
        all_variables = all_variables.set_index(all_variables['variable'])
        all_variables = all_variables.query(
            "(status == 'modelisable' or status == 'utilisable') and (type == 'categoriel' or (type =='numerique' and sous_type=='float'))")

        corr_df = pd.DataFrame(np.zeros((all_variables.shape[0], all_variables.shape[0])),index=all_variables['variable'], columns=all_variables['variable'])

        for num_var1, name_var1 in enumerate(all_variables['variable']):

            type_var1 = all_variables['type'][name_var1]

            if (type_var1 == 'categoriel'):
                if (df[name_var1].isnull().apply(lambda x : 1 if x else 0).sum()>0 ):
                    df[name_var1] = df[name_var1].cat.add_categories(['aaValuezz'])
                    df[name_var1] = df[name_var1].fillna(value = 'aaValuezz')


            for num_var2, name_var2 in enumerate(all_variables['variable']):
                if num_var2 <= num_var1:
                    continue
                print(str(num_var1)+':'+str(num_var2)+' ; '+ name_var1+':'+name_var2)
                type_var2 = all_variables['type'][name_var2]
                if (type_var2 == 'categoriel'):
                    if (df[name_var2].isnull().apply(lambda x : 1 if x else 0).sum()>0 ):
                        df[name_var2] = df[name_var2].cat.add_categories(['aaValuezz'])
                        df[name_var2] = df[name_var2].fillna(value = 'aaValuezz')

                model = RandomForestClassifier(n_estimators=100)
                result = getRFcorrelation(df, name_var1, name_var2, model = model)


                corr_df[name_var1][num_var2] = result
                corr_df[name_var2][num_var1] = result

        # Handle the errors when there is only one modality in a variable
        corr_df = corr_df.fillna(value = 0)
        for num_var1, name_var1 in enumerate(all_variables['variable']):
            corr_df[name_var1][num_var1] = 0.0

        # Set up the matplotlib figure
        from matplotlib import rcParams
        rcParams.update({'figure.autolayout': True})
        fig, ax = plt.subplots()
        name = "temp.png"

        # Draw the heatmap using seaborn
        #heatmap = sns.heatmap(corr_df, vmax=1.0, square=True, cmap=sns.diverging_palette(10, 220, as_cmap=True))
        heatmap = sns.heatmap(corr_df, vmax=1.0, square=True, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))
        plt.xticks(rotation = 'vertical')
        plt.yticks(rotation = 'horizontal')
        heatmap.get_figure().savefig(name)
        plt.gcf().clear()
        plt.close(fig)
        document.add_picture(name, width=Cm(17.0), height=Cm(22.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


        # linkage = sp.cluster.hierarchy.linkage(sp.spatial.distance.squareform(1-corr_df), method='average')
        # fig, ax = plt.subplots()
        # heatmap_clustered = sns.clustermap(corr_df, row_linkage=linkage, col_linkage=linkage, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))
        # # heatmap_clustered = sns.clustermap(corr_df, cmap=sns.light_palette((210, 90, 60), input="husl", as_cmap = True))

        # name = "temp.png"
        # plt.setp(heatmap_clustered.ax_heatmap.get_yticklabels(), rotation=0)
        # plt.setp(heatmap_clustered.ax_heatmap.get_xticklabels(), rotation=90)
        # heatmap_clustered.savefig(name)
        # document.add_picture(name, width=Cm(17.0), height=Cm(22.0))
        # last_paragraph = document.paragraphs[-1]
        # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        return document

class AddonManager:
    def __init__(self, addons, field_addons):
        self.dict_addons = []
        self.dict_field_addons = {}
        for addon in addons:
            self.dict_addons= self.dict_addons + [addon]
        for addon in field_addons:
            self.dict_field_addons[addon.field] = self.dict_field_addons.get(addon.field, []) +[addon]

    def __call__(self):
        return self.dict_addons, self.dict_field_addons


if __name__=='__main__':
    example = ExampleFieldAddon('test', 'f1')
    example2 = ExampleFieldAddon('test1', 'f2')
    example3 = ExampleFieldAddon('test', 'f1')

    example4 = Addon('test1')
    manager = AddonManager([example4], [example, example2, example3, ])

    print(manager())
