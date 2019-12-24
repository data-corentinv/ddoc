"""
This class allows the generation of a report given a directory
It can be used for two things : Explore raw sources or generate a final report
on all the sources used
"""

from docx import Document
from docx.shared import Cm, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from numbers import Number
import pandas as pd
import numpy as np
from datetime import datetime
import ntpath
# Plot
import matplotlib as mpl
import matplotlib.pyplot as plt
import seaborn as sns
from pdb import set_trace as pause

from .report_generator import ReportGenerator

# Project dependencies, to remove if you want reusability

cwd = os.path.dirname(os.path.realpath(__file__))

config = {
    'data_paths': {
        'store_dir' : ''
    },
    "max_categories_plot" : 100,
    "alerte" : {
        'réfléchir' : os.path.join(cwd, 'assets/arevoir.png'),
        "utilisable" : os.path.join(cwd, 'assets/utilisable.png'),
        "inutilisable" : os.path.join(cwd, 'assets/inutilisable.png'),
        "modelisable" : os.path.join(cwd, 'assets/modelisable.png')
    }
}

styles = {

   "warning" : {
        "color" : RGBColor(0xFF, 0x00, 0x00)
    },
    "note" : {
        "italic" : True
    }
}


def is_categoric(series, treshold=50 ):
    return len(series.unique()) < treshold

def is_numeric(series):
    return np.issubdtype(series.dtype, np.number)

def plot(df, column, is_categoric=False, is_numeric=False, is_integer=False, dims=None):
    df_plot = df
    # if is_categoric and len(df_plot[column].unique())> config['max_categories_plot']:
    #     top_n = df[column].value_counts().reset_index()['index'].values
    #     top_n = top_n[:min(top_n.shape[0], config['max_categories_plot'])]
    #     df_plot = df_plot[df_plot[column].isin(top_n)]
    fig, ax = plt.subplots(figsize=dims)
    name = "temp.png"
    # pause()
    if is_categoric and len(df_plot[column].unique())<= config['max_categories_plot']:
        order = df_plot[column].value_counts().reset_index()['index']
        # Filter labels name to keep only first 20 caracters
        order = order.map(str)
        #order = order.str[:10]
        order = order.values
        print(order)

        A = sns.countplot(ax=ax, x=column,order=order, data=df_plot).get_figure()
        plt.xticks(rotation=90)
        try:
            plt.tight_layout()
        except:
            plt.xticks(rotation=-90)
        A.savefig(name)
        plt.gcf().clear()
    elif is_numeric:
        if len(df_plot[column].value_counts())<4:
            to_plot = df_plot[df_plot[column].notnull()][column]
            nb_outliers_up = 0
            nb_outliers_down = 0
        else:
            mean = df_plot[column].mean()
            std = df_plot[column].std()
            mask_up = (df_plot[column] < mean + 3*std) # Replace with quantiles ?
            mask_down= (df_plot[column] > mean - 3*std) # Replace with quantiles ?
            mask = mask_up & mask_down
            to_plot = df_plot[mask]
            to_plot = to_plot[to_plot[column].notnull()]
            to_plot = to_plot[column]
            nb_outliers_up = df_plot[~mask_up].shape[0]
            nb_outliers_down = df_plot[~mask_down].shape[0]
        if not is_integer: # We first suppose it is countinuous
            # On coupe à 3STD
            A = sns.distplot(to_plot, hist_kws={'rwidth' : 0.9}, color=sns.xkcd_rgb["dusty blue"], ax=ax, norm_hist=True, kde=False).get_figure()
            plt.text(0.5, 0.5,'Outliers up:'+str(nb_outliers_up)+' Outliers down : '+str(nb_outliers_down), horizontalalignment='center',  verticalalignment='center', transform=ax.transAxes)
            plt.tight_layout()
            A.savefig(name)
            plt.gcf().clear()
        else: #entier
            binwidth = 1
            # print(column)
            # pause()
            data_plot = to_plot.astype(int)
            bins = np.arange(min(data_plot)-binwidth/2, max(data_plot) + binwidth*1.5, binwidth)
            A = sns.distplot(data_plot, bins = bins, hist_kws={'rwidth' : 0.9}, color=sns.xkcd_rgb["dusty blue"], ax=ax, norm_hist=False, kde=False).get_figure()
            plt.text(0.5, 0.5,'Outliers up:'+str(nb_outliers_up)+' Outliers down : '+str(nb_outliers_down), horizontalalignment='center',  verticalalignment='center', transform=ax.transAxes)
            plt.tight_layout()
            A.savefig(name)
            plt.gcf().clear()
    else:
        #use treemap for non numeric and too many categories
        return None
    plt.close(fig)
    return name

class WordGenerator(ReportGenerator):
    """
     This class is providing a tool to generate a report on json files automatically
     You just have to specify the folder the ReportGenerator will have to work with
     then just call ReportGenerator() to generate your report
    """
    def __init__(self, table=None,out_directory='', metadata_location=None, field_addons=[], addons=[]):
        """
        :param directory: directory the ReportGenerator will work on
        :param sections_filters: It is a dict with the format {title : filter, title2 : filter2 }
        filter are filtering function (You get back a boolean for each file)
        """
        super().__init__(table=table, out_directory=out_directory, metadata_location=metadata_location, field_addons=field_addons, addons=addons)

    def configurate(self, document, margin=2):
        # changing the page margins
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(margin)
            section.bottom_margin = Cm(margin)
            section.left_margin = Cm(margin)
            section.right_margin = Cm(margin)

        # Configurate headings #78B62A #700D35 #112B29
        #lmf_color = RGBColor(0x11, 0x2b, 0x29)
        #lmf_color = RGBColor(0xb4, 0x52, 0xcd)
        roquette_color = RGBColor(0,0,0)
        style = document.styles['Title']
        style.font.color.rgb = roquette_color

        style = document.styles['Heading 1']
        style.font.color.rgb = roquette_color
        style.font.size = Pt(16)

        style = document.styles['Heading 2']
        style.font.color.rgb = roquette_color
        style.font.size = Pt(14)

        style = document.styles['Heading 3']
        style.font.color.rgb = roquette_color
        style.font.size = Pt(12)

        style = document.styles['Normal']
        font = style.font
        font.name = 'Roboto'#Calibri
        font.size = Pt(10)

    def generate(self, **kwargs):
        """
        For each data file in directory create a type description
        """
        document = Document()
        self.configurate(document)
        document.add_heading('Dictionnaire de données', 0)
        df = self.load_df(**kwargs)
        self.get_types_report(df, document)
        document.save(os.path.join(self.out_directory, 'type_report'+ '.docx'))

    def describe(self, document, df, column, metadata):
        series = df[column]
        values, values_raw, counts_raw = self.get_values(series)
        if(len(str(values)) > 1000):
            values = "Examples are too long to be printed"
        p = document.add_heading(str(column), 3)
        #document.add_picture(config.get('alerte', {}).get(metadata.get('alerte',''), config.get('alerte').get('réfléchir')), width=Cm(2.0), height=Cm(2.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if 'description' in metadata:
            p = document.add_paragraph()
            p = p.add_run(metadata.get('description', ''))

        if 'warning' in metadata:
            p = document.add_paragraph()
            r = self.write(p, metadata.get('warning', ''), 'warning')
        if 'note' in metadata:
            p = document.add_paragraph()
            r = self.write(p, metadata.get('note', ''), 'note')

        # Plot
        try:
            is_categoric_ = metadata.get('type','categoriel') == 'categoriel'
            is_numeric = (metadata.get('type', 'inconnu') == 'numerique')
            is_integer = (metadata.get('sous-type', 'inconnu') == 'integer')
            file_to_plot = plot(df, column, is_categoric=is_categoric_,is_numeric=is_numeric, is_integer = is_integer,dims=(8,4))
            if column == 'TARGET': #on ne veut pas afficher les proportions de la target
                file_to_plot=None
            if file_to_plot is not None:
                document.add_picture(file_to_plot, width=Cm(18.0), height=Cm(9.0))
                last_paragraph = document.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print("Error "+ str(e))

        document.add_paragraph()

        table = document.add_table(rows=2, cols=5)
        table.autofit = True
        table.style = 'LightGrid-Accent1'
        table.cell(0, 0).text = 'Valeurs manquantes'
        table.cell(0, 1).text = 'Type'
        table.cell(0, 2).text = 'Sous-Type'
        table.cell(0, 3).text = 'Nombre de valeurs uniques'
        table.cell(0, 4).text = 'Exemples de valeurs'

        r = table.cell(1, 0).paragraphs[0].add_run("{:d} ({:.2f}".format(series.isnull().sum(), 100*series.isnull().sum()/series.shape[0]) + '%)')
        r.bold = False
        table.cell(1, 1).text = metadata.get('type', 'inconnu')
        table.cell(1, 2).text  = metadata.get('sous-type', self.get_type(df, column))
        table.cell(1, 3).text = str(len(series[series.notnull()].unique()))

        table.cell(1, 4).text = str(values)


        # New values
        document.add_paragraph()
        table = document.add_table(rows=2, cols=4)
        table.autofit = True
        table.style = 'LightGrid-Accent1'
        table.cell(0, 0).text = 'Valeur min'
        table.cell(0, 1).text  = 'Valeur max'
        table.cell(0, 2).text = 'Provenance'
        table.cell(0, 3).text = 'Origine'
        if series.dtype.kind in 'bifc' and not metadata.get('type', False)=='categoriel':
            r = table.cell(1, 0).paragraphs[0].add_run(str(series.min()))
            r.bold = False
            table.cell(1, 1).text  = str(series.max())
        else:

            r = table.cell(1, 0).paragraphs[0].add_run(str(series[series.notnull()].map(lambda x: len(str(x))).astype(np.float).min()))
            r.bold = False
            table.cell(1, 1).text  = str(series[series.notnull()].map(lambda x: len(str(x))).astype(np.float).max() )
            table.cell(0, 0).text = 'Taille minimale'
            table.cell(0, 1).text  = 'Taille maximale'
        table.cell(1, 2).text = metadata.get('provenance', '')
        table.cell(1, 3).text = metadata.get('origine', '')

    def write(self, paragraph, text, style=None):
        p = paragraph.add_run(text)
        if style is not None:
            p.font.color.rgb = styles.get(style, {}).get('color', RGBColor(0x00, 0x00, 0x00))
            p.italic = styles.get(style, {}).get('italic', False)
        return p

    def get_values(self, series):
        values_raw = series.value_counts().reset_index()['index'].values
        counts_raw = series.value_counts().values
        values_example = values_raw.astype(str)[0:min(len(values_raw), 5)]
        return values_example, values_raw, counts_raw

    def document_field(self, document, df, column, metadata, max_value_categorical=50):
        """
        Document one field
        """
        series = df[column]
        print("Dealing with "+str(column))
        _ , values_raw, counts_raw = self.get_values(series)

        self.describe(document, df, column, metadata)
        # Add dictionnary
        p = document.add_paragraph()
        if is_categoric(series,treshold=max_value_categorical):
            table = document.add_table(rows=series.value_counts().shape[0]+1, cols=4)
            table.style = 'LightGrid-Accent1'
            #table.style = 'TableGrid'
            table.cell(0,0).text = 'Valeur'
            table.cell(0,1).text = 'Taux'
            table.cell(0,2).text = 'Nombre'
            table.cell(0,3).text = 'Description'
            metadata_ = dict(metadata.get('valeurs', {}))

            for idx, index_raw in enumerate(zip(values_raw, counts_raw)):
                value, count_raw = index_raw
                table.cell(idx+1, 1).text = "{:.2f}%".format(100*count_raw/df.shape[0])
                table.cell(idx+1, 2).text = str(count_raw)
                try:
                    str_value = str(value)
                    if isinstance(value, Number):
                        if str(value).endswith('.'):
                            str_value=str_value[:-1]
                        elif str(value).endswith('.0'):
                            str_value=str_value[:-2]
                    table.cell(idx+1, 0).text = str_value
                except:
                    table.cell(idx+1, 0).text = "? Encoding Issue ?"
                # The issue is that a key can be an integer or a float that can ends with ".", make it generic for json comparison
                str_value = str(value)
                if isinstance(value, Number):
                    if str(value).endswith('.'):
                        str_value=str_value[:-1]
                    elif str(value).endswith('.0'):
                        str_value=str_value[:-2]
                table.cell(idx+1, 3).text = metadata_.pop(str_value, "")
            for key, item in metadata_.items():
                row = table.add_row()
                row.cells[0].text = key
                row.cells[1].text = "Non présent"
                row.cells[2].text = '0'
                row.cells[3].text = item

        # Add addons
        for addon in self.field_addons:
            document = addon(document, metadata, df, column)
            plt.close('all')

        document.add_paragraph()


    def get_types_report(self, df, document,  max_value_categorical=50, **kwargs):
        """
        Generate fields descriptors of a table
        """
        # Add addons
        #addons_table, field_addons = self.addons_manager() # deprecated
        field_addons = self.field_addons
        self.current_table_field_addons = field_addons
        metadata = self.metadata
        #document.add_picture('/home2/corentinV/myPackages/datadoc_tmp/data/logo/leroy_merlin_logo.jpg', width=Cm(4.0), height=Cm(2.0))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_heading('Description générale', 2)
        document.add_paragraph(metadata.get('description', ""))
        table_ = document.add_table(rows=2, cols=2)
        table_.style = 'LightGrid'
        table_.cell(0, 0).text = 'Nombre de lignes'
        table_.cell(0, 1).text = 'Nombre de colonnes'
        table_.cell(1, 0).text = str(df.shape[0])
        table_.cell(1, 1).text = str(len(df.columns))
        document.add_heading('Champs', 2)
        metadata = metadata.get('champs', {})
        for column in sorted(df.columns):
            self.document_field(document, df, column, metadata.get(str(column), {}), max_value_categorical=max_value_categorical)
            document.add_paragraph()
            document.add_page_break()

        # Add general addons
        if self.addons != "none":
            for addon in self.addons:
                document = addon(document, metadata, df)
            document.add_page_break()
