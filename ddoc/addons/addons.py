from matplotlib import pyplot as plt
import pandas as pd
from docx.shared import Cm

#import sys
#sys.path.append('../')
from ..reporting.utils_for_supervized_learning import *
from .utils.utils import study_feature_reconstruct


class AddonIsTarget:
    def __call__(self, document, metadata, df, field):
        study_feature_reconstruct(df, field, label='TARGET')
        plt.savefig('temp2.png')
        document.add_picture('temp2.png', width=Cm(15.0), height=Cm(9.0))
        return document


class AddonIsTargetQuantile:
    def __call__(self, document, metadata, df, field ):
        target ='TARGET'

        color_rectangle = sns.xkcd_rgb["dusty green"]
        color_rectangle_nan = sns.xkcd_rgb["dusty red"]
        axe_fontsize = 10
        proportion_minimum_to_display_unique = 0.03

        if df[field].isnull().sum() == len(df): # only NaN
            df[field] = pd.Categorical(df[field].fillna("NaN"), categories = ['NaN'])


        if (field != target) and (df[field].dtype == float or df[field].dtype==np.float32):
            # case where numerical
            number_of_quantiles = 100.0

            plot_numerical_variable(df, target, field, number_of_quantiles, color_rectangle, color_rectangle_nan, axe_fontsize, proportion_minimum_to_display_unique)
            document.add_picture('temp3.png', width=Cm(18.0), height=Cm(9.0))

        elif field != target:
            # case where categorical

            plot_categorical_variable(df, target, field, color_rectangle, color_rectangle_nan, axe_fontsize, proportion_minimum_to_display_unique)
            document.add_picture('temp3.png', width=Cm(18.0), height=Cm(9.0))
        else :
            #
            print(field + ' is the target.')

        return document
